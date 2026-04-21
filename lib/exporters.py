from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_wrong_questions_docx(
    nickname: str,
    questions: list[dict],
    question_states: dict[str, dict],
) -> BytesIO:
    document = Document()
    document.core_properties.author = "Easy-Request"
    document.core_properties.title = "医学伦理学错题导出"

    title = document.add_heading("医学伦理学错题复盘", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary = document.add_paragraph()
    summary.add_run(f"昵称：{nickname}\n").bold = True
    summary.add_run(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    summary.add_run(f"错题数量：{len(questions)}")

    document.add_paragraph(" ")

    for index, question in enumerate(questions, start=1):
        state = question_states.get(question["question_id"], {})
        document.add_heading(f"{index}. {question['stem']}", level=1)

        meta = document.add_paragraph()
        meta.add_run("来源：").bold = True
        meta.add_run(
            f"{question['source_doc']} / {question['section_title']} / {question['question_type']}"
        )
        if question.get("group_range"):
            meta.add_run(f" / 共用备选范围 {question['group_range']}")

        for option in question["options"]:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(f"{option['key']}. ").bold = True
            paragraph.add_run(option["text"])

        answer = document.add_paragraph()
        answer.add_run("正确答案：").bold = True
        answer.add_run(question["answer"])

        wrong = document.add_paragraph()
        wrong.add_run("最近一次错误答案：").bold = True
        wrong.add_run(state.get("last_wrong_selected") or "暂无记录")

        mastery = document.add_paragraph()
        mastery.add_run("当前状态：").bold = True
        mastery.add_run("已掌握" if state.get("mastered") else "待巩固")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
